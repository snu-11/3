import streamlit as st
import koreanize_matplotlib 
import os
import koreanize_matplotlib 
from docx import Document
from io import BytesIO
import pandas as pd

# 워드 문서로 변환하는 함수
def convert_to_word(dataframe):
    doc = Document()
    for index, row in dataframe.iterrows():
        doc.add_heading('생각1', level=1)
        student_thought = str(row['생각1: 뉴스 내용과 AI 노출지수에 대한 의견']) if pd.notna(row['생각1: 뉴스 내용과 AI 노출지수에 대한 의견']) else ""
        doc.add_paragraph(student_thought)

        doc.add_heading('생각2', level=1)
        job_relation_thought = str(row['생각2: 어떤 직업이 사라질 것 같은가']) if pd.notna(row['생각2: 어떤 직업이 사라질 것 같은가']) else ""
        doc.add_paragraph(job_relation_thought)

        doc.add_page_break()
    return doc

# Streamlit 페이지 설정
st.set_page_config(
    page_title="AI 노출지수와 직업",
    page_icon="🤖",
    layout="centered",
    initial_sidebar_state="auto",
    menu_items=None
)

# Streamlit 앱 제목 설정
st.markdown("<h1 style='margin-bottom: 0;'> AI 기술로 수행 가능한 업무는?</h1>", unsafe_allow_html=True)

st.divider()
st.video('https://www.youtube.com/watch?v=6XXJZh7U1xk', format="video/mp4", start_time=0)

# 직업 데이터 및 관련 이모지
job_data = {
    '직업': ['의사', '교사', '개발자', '회계사', '변호사', '가수', '요리사',
            '디자이너', '프로게이머', '운동선수'],
    'AI 노출지수': [93, 1, 94, 81, 79, 0, 6, 13, 60, 0],
    '이모지': ['👨‍⚕️', '👩‍🏫', '💻', '🧾', '⚖️', '🎤', '👨‍🍳', 
              '🎨', '🎮', '🏃‍♂️']
}

st.divider()

# Streamlit 앱 제목 설정
st.markdown("<h1 style='margin-bottom: 0;'>직업별 AI 노출지수 추측 게임</h1>", unsafe_allow_html=True)

# 점수를 위한 세션 상태 초기화
if 'score' not in st.session_state:
    st.session_state.score = 0

# 제출 버튼 상태 초기화
if 'submitted' not in st.session_state:
    st.session_state.submitted = False

# 각 직업별 추측 입력 및 즉각적인 피드백
for i in range(0, len(job_data['직업']), 5):
    cols = st.columns(5)
    for j in range(5):
        if i + j < len(job_data['직업']):
            job = job_data['직업'][i + j]
            emoji = job_data['이모지'][i + j]
            correct_answer = job_data['AI 노출지수'][i + j]

            with cols[j]:
                # HTML 스타일을 사용하여 글자 크기 조정
                st.markdown(f"<span style='font-size: 15px;'>{emoji} {job}</span>", unsafe_allow_html=True)
                
                # Unique key for each number input
                guess_key = f"guess_{job}"
                guess = st.number_input("", min_value=0, max_value=100, step=1, key=guess_key)

                # 제출 버튼이 눌렸다면, 피드백 제공
                if st.session_state.submitted:
                    if abs(guess - correct_answer) <= 5:
                        st.success(f"🎉 정답! AI 노출지수는 {correct_answer}입니다.")
                        if st.session_state.score <= i + j * 10:
                            st.session_state.score += 10  # 점수 추가
                    else:
                        st.error(f"❌ 틀렸습니다. 정답은 {correct_answer}입니다.")
                        if st.session_state.score >= (i + j) * 10 + 10:
                            st.session_state.score -= 10  # 점수 제거

    # 각 세션 사이에 두꺼운 수평선 추가
    st.markdown("<hr style='height:0.5px;border-width:0;color:gray;background-color:gray'>", unsafe_allow_html=True)

# 제출 버튼
st.session_state.submitted = st.button('제출')

# 총점 표시
st.markdown(f"## 🏆 총점: {st.session_state.score}점")


st.divider()

# Streamlit captions and subheaders for the questions
st.caption(":blue_heart:비판적사고 기르기!")
st.subheader("위 뉴스 내용과 AI 노출지수 값에 동의하나요?")

# Unique key for the first text area
first_student_thought_key = "first_student_thought"
first_student_thought = st.text_area("나의 의견을 적어주세요🖊️", key=first_student_thought_key)

st.subheader("고용 현황 및 AI 노출 지수 내용을 바탕으로 어떤 직업이 사라질 것 같나요?")

# Unique key for the second text area
second_student_thought_key = "second_student_thought"
second_student_thought = st.text_area("나의 의견을 적어주세요🖊️", key=second_student_thought_key)

# Check if the submission button is pressed
if st.button("제출", key="final_submit"):
    # Create or load the existing DataFrame
    if 'student_thoughts.csv' not in os.listdir():
        student_thoughts_df = pd.DataFrame(columns=['생각1: 뉴스 내용과 AI 노출지수에 대한 의견', '생각2: 어떤 직업이 사라질 것 같은가'])
    else:
        student_thoughts_df = pd.read_csv('student_thoughts.csv', encoding='utf-8')

    # Print or log the column names for debugging
    print("Column names in the DataFrame:", student_thoughts_df.columns)

    # Append new data to the DataFrame
    new_data = pd.DataFrame({'생각1: 뉴스 내용과 AI 노출지수에 대한 의견': [first_student_thought],
                             '생각2: 어떤 직업이 사라질 것 같은가': [second_student_thought]})
    student_thoughts_df = pd.concat([student_thoughts_df, new_data], ignore_index=True)
    student_thoughts_df.to_csv('student_thoughts.csv', index=False, encoding='utf-8')

    # Call the convert_to_word function
    doc = convert_to_word(student_thoughts_df)

    # Display the submitted data
    st.subheader("나의 생각:")
    st.write("뉴스 내용과 AI 노출지수에 대한 의견:", first_student_thought)
    st.write("어떤 직업이 사라질 것 같나요:", second_student_thought)

    # Convert to Word document and provide download link
    doc = convert_to_word(student_thoughts_df)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(label="워드 문서로 다운로드",
                       data=buffer,
                       file_name="student_thoughts.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
