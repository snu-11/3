import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
from matplotlib.ticker import MaxNLocator
import seaborn as sns 
import koreanize_matplotlib 
from docx import Document
from io import BytesIO

# 워드 문서로 변환하는 함수
def convert_to_word(dataframe):
    doc = Document()
    for index, row in dataframe.iterrows():
        doc.add_heading('학생 생각', level=1)
        # Convert to string before adding to the document
        student_thought = str(row['학생 생각']) if pd.notna(row['학생 생각']) else ""
        doc.add_paragraph(student_thought)

        doc.add_heading('직업과의 연관성 생각', level=1)
        # Convert to string before adding to the document
        job_relation_thought = str(row['직업과의 연관성 생각']) if pd.notna(row['직업과의 연관성 생각']) else ""
        doc.add_paragraph(job_relation_thought)

        doc.add_page_break()
    return doc


# 한글 폰트 설정
plt.rcParams['font.family'] = 'NanumGothic'
plt.rcParams['axes.unicode_minus'] = False

# Streamlit 페이지 설정
st.set_page_config(
    page_title="나에게 맞는 직업",
    page_icon="📈",
    layout="centered",
    initial_sidebar_state="auto"
)

st.title("나에게 맞는 직업찾기🔍")
st.subheader(":gray[나에게 중요한 가치는 무엇일까?]")
st.write("직업가치관 검사 결과를 다시 확인하려면 [여기](https://www.career.go.kr)를 클릭하세요.")


st.divider()

st.subheader("나의 가치에 맞는 직업을 찾기 위한 선그래프 그리기📈")


# 엑셀 파일 업로드
uploaded_file = st.file_uploader("엑셀 파일을 업로드하세요.", type=["xlsx", "xls"])

# 엑셀 파일 불러오기 및 데이터 확인
if uploaded_file is not None:
    df = pd.read_excel(uploaded_file, header=0, index_col=0)

    # 데이터의 일부 출력
    st.write("데이터 미리보기:", df.head())

    st.success('파일 업로드 성공!')

    # 행과 열 선택 옵션
    row_options = df.index.tolist()
    column_options = df.columns.tolist()

    selected_rows = st.multiselect("행 선택", options=row_options)
    selected_columns = st.multiselect("열 선택", options=column_options)

    if selected_rows and selected_columns:
        # 선택된 행과 열에 대한 데이터프레임
        selected_df = df.loc[selected_rows, selected_columns]

        # 데이터를 정수형으로 변환
        selected_df = selected_df.apply(pd.to_numeric, errors='coerce').fillna(0).astype(int)

        # 데이터 시각화
        st.subheader("선택한 데이터 시각화")
        fig, ax = plt.subplots()
        selected_df.plot(ax=ax, kind='line', marker='o')
        st.pyplot(fig)


st.divider()


# 사용자 입력
first_student_thought = st.text_area("예) 나는 **돈**이 중요해!, 나는 **워라밸**이 중요해!")
second_student_thought = st.text_area("그래프를 통해 발견한 내용을 적어주세요🖊️")

if st.button("제출"):
    # 생각을 DataFrame에 저장
    if 'student_thoughts.csv' not in os.listdir():
        student_thoughts_df = pd.DataFrame(columns=['학생 생각', '직업과의 연관성 생각'])
    else:
        student_thoughts_df = pd.read_csv('student_thoughts.csv', encoding='utf-8')

    new_data = {'학생 생각': first_student_thought, '직업과의 연관성 생각': second_student_thought}
    student_thoughts_df = student_thoughts_df.append(new_data, ignore_index=True)
    student_thoughts_df.to_csv('student_thoughts.csv', index=False, encoding='utf-8')

    # 제출한 데이터를 화면에 표시
    st.subheader("나의 생각:")
    st.write("중요한 가치:", first_student_thought)
    st.write("직업과의 연관성:", second_student_thought)

    # 워드 문서로 변환 및 다운로드 링크 제공
    doc = convert_to_word(student_thoughts_df)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(label="워드 문서로 다운로드",
                       data=buffer,
                       file_name="student_thoughts.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


